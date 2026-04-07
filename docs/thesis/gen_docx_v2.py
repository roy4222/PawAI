"""Generate 114-thesis.docx matching 106 SA template format exactly."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微軟正黑體'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    if level == 1:
        hs.font.size = Pt(16)
    elif level == 2:
        hs.font.size = Pt(14)

# ── Helpers ──
def merge_cells(table, row, col_start, col_end):
    """Merge cells in a row."""
    cell = table.cell(row, col_start)
    for c in range(col_start + 1, col_end + 1):
        cell.merge(table.cell(row, c))
    return cell

def add_story_card(name, category, role, need, value, acceptance_items, priority=None):
    """User Story card in 106 template format: 4-col table."""
    # Row count: header(1) + content(1) + optional priority(1)
    nrows = 3 if priority else 2
    table = doc.add_table(rows=nrows, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: 名稱 | name | 類別 | category
    table.cell(0, 0).text = '名稱'
    table.cell(0, 1).text = name
    table.cell(0, 2).text = '類別'
    table.cell(0, 3).text = category
    for c in [0, 2]:
        for p in table.cell(0, c).paragraphs:
            for r in p.runs:
                r.bold = True

    # Row 1: merged 4 cols → 角色/需求/價值/接受條件
    cell = merge_cells(table, 1, 0, 3)
    cell.text = ''
    p = cell.paragraphs[0]
    p.add_run(f'角色：{role}\n').bold = False
    p.add_run(f'需求：{need}\n').bold = False
    p.add_run(f'價值：{value}\n').bold = False
    p.add_run('\n接受條件：\n').bold = True
    for item in acceptance_items:
        p.add_run(f'● {item}\n').bold = False

    # Row 2: priority
    if priority:
        table.cell(2, 0).text = '優先順序'
        for p in table.cell(2, 0).paragraphs:
            for r in p.runs:
                r.bold = True
        merge_cells(table, 2, 1, 3).text = str(priority)

    doc.add_paragraph()

def add_data_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()

def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p
def bullet(text): doc.add_paragraph(text, style='List Bullet')
def page_break(): doc.add_page_break()

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('114專題文件'); run.font.size = Pt(26); run.bold = True
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Paw AI 機器狗'); run.font.size = Pt(30); run.bold = True
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing = 2.0
for line in ['指導老師：董惟鳳', '412401123 楊沛蓁', '412401355 盧柏宇', '412401135 鄔雨彤', '412401082 陳如恩', '412401094 黃  旭']:
    p.add_run(line + '\n')

# ══════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════
page_break()
h1('第一章 系統描述')
h2('一、發展背景與動機')
para('研究背景', bold=True)
para('近年來，人工智慧技術發展迅速，特別是大型語言模型（Large Language Models, LLMs，如 Qwen、LLaMA）的突破，使得機器系統具備了深度的語意理解與邏輯推理能力。同時，邊緣運算與電腦視覺（如 MediaPipe 姿態辨識、YuNet 人臉偵測技術）以及語音處理技術（如 SenseVoice 語音辨識與 edge-tts 語音合成）的日益成熟，讓資訊系統能夠以多模態（Multi-modal）的方式，精準感知周遭環境與人類的真實意圖。')
para('在終端硬體載具方面，四足機器人（如仿生機器狗）的機電控制與運動靈活性已有顯著的提升。然而，綜觀目前市面上的實體機器人應用，多半仍依賴傳統的遙控設備、手機 App 或是僵化的預設腳本進行操作。現有系統缺乏將「雲端高階 AI 大腦」與「邊緣終端硬體（ROS2）」深度結合的解決方案，導致人機互動往往流於單向、死板的指令接收，機器無法根據人類自然的肢體語言（手勢）、語音指令或當下的環境脈絡，進行即時的動態感知與自主決策。')

para('研究動機', bold=True)
para('基於上述背景，本專題「PawAI」旨在打破傳統實體機器人僵化且高門檻的互動模式。我們運用資訊管理領域的系統整合（System Integration）與軟體工程技術，開發一套無縫串接視覺感知、語音互動與大型語言模型決策的智能控制系統。')
para('本研究的核心動機包含以下幾點：')
para('1. 實現自然且直覺的人機互動：透過建構完整的多模態感知管線（Multi-modal Perception Pipeline），讓系統能即時捕捉並解析使用者的手勢與臉部特徵；結合語音雙向處理，讓使用者能以最自然的「說話」與「動作」與機器狗互動，免除繁瑣的傳統控制介面。')
para('2. 賦予終端設備動態決策能力：藉由串接大型語言模型 Qwen2.5-7B-Instruct（透過 vLLM 部署於遠端 RTX 8000 伺服器）作為系統的核心邏輯中樞，將接收到的多模態感知資料轉化為結構化的意圖（Intent），讓機器狗不再只是單純執行寫死的程式碼，而是能「理解」複雜情境並規劃後續行為。')
para('3. 軟硬體整合與控制迴路優化：透過建構穩定的 ROS2 節點通訊與統一中控模組（Interaction Executive），確保 AI 模型的決策結果能以極低的延遲、安全且精確地轉化為實體機器狗的具體動作，完成從「感知」、「思考」到「行動」的完整閉環。')
para('總結而言，本專題期望透過 PawAI 系統的建置，探索大型語言模型結合實體機器人的無限可能，不僅為智慧陪伴、巡檢或展演等應用場景提供具高度擴充性的架構範例，更進一步實現真正意義上的智慧人機共作（Human-Robot Collaboration）。')

# ── 二、系統發展目的 ──
h2('二、系統發展目的')
add_data_table(['問題', '解決方案'], [
    ['對互動使用者而言，操作實體機器人需透過遙控器或手機 App，互動方式單向且不直覺。', '整合 MediaPipe 手勢/姿勢辨識、YuNet 人臉辨識與語音模組，實現自然直覺的多模態互動。'],
    ['機器人僅能執行預先寫死的腳本，缺乏情境理解與自主決策能力。', '串接大語言模型（Qwen2.5-7B-Instruct）作為核心大腦，依據當下脈絡動態生成行為回饋。'],
    ['雲端 AI 模型與邊緣硬體設備之間，缺乏穩定、低延遲的整合與執行機制。', '開發基於 ROS2 的統一中控模組（Interaction Executive），精準且安全地驅動終端硬體作動。'],
    ['對操作者而言，傳統機器人系統難以遠端監控與除錯。', '建立 PawAI Studio 網頁控制台，提供即時感知面板、語音/文字互動與即時影像串流。'],
])

para('(一) 整合多模態感知技術，提供自然直覺的人機互動體驗', bold=True)
para('打破傳統單向且受限的控制介面，本系統導入 MediaPipe 模型進行即時的姿態與手勢推論，並輔以 YuNet + SFace 人臉辨識技術鎖定互動對象。同時，結合 STT（語音轉文字）與 TTS（文字轉語音）模組，讓使用者能以最日常的肢體動作與口語對話，與系統進行雙向、流暢的互動。')
para('(二) 串接大語言模型，建構具備情境理解的自主決策中樞', bold=True)
para('為賦予機器狗動態決策的能力，本專案將前端擷取的多模態感知資料轉化為結構化的意圖（Intent），並透過 HTTP API 串接部署於遠端伺服器的大語言模型（Qwen2.5-7B-Instruct）。藉此，系統能深入理解使用者的語意與當下的互動脈絡，取代傳統預先寫死的控制腳本，自主規劃並生成最合適的行為回饋。')
para('(三) 開發 ROS2 統一中控，實現低延遲且穩定的硬體作動閉環', bold=True)
para('為確保雲端 AI 的決策能精確落實於邊緣實體設備，本系統建構了基於 ROS2 的節點通訊機制。透過統一中控模組 Interaction Executive（以 State Machine 實現事件路由與優先序仲裁），系統能將抽象的行為指令，以極低的延遲轉譯為機器狗能直接執行的動作控制訊號。')
para('(四) 建立 PawAI Studio 網頁控制台，確保系統的高可觀測性與可維護性', bold=True)
para('建立 PawAI Studio 作為系統的統一觀測與操控入口，整合即時感知面板、語音/文字互動介面、以及即時影像串流，方便團隊進行效能監控與錯誤排查。')

para('非功能性需求', bold=True)
bullet('低延遲：語音端到端互動延遲目標 < 5 秒')
bullet('高可用性：三級降級策略（雲端 → 本地 → 規則），確保無網路時仍可基本互動')
bullet('模組化：每個感知模組為獨立 ROS2 package，可單獨替換或升級')
bullet('安全性：統一中控模組仲裁所有動作指令，緊急事件（跌倒）具最高優先序')

# ── 三、系統範圍 ──
h2('三、系統範圍')
para('本節根據使用者的互動情境，盤點系統的功能性需求。')

for title, items in [
    ('1. 語音互動模組', ['語音辨識（ASR）：三級自動降級（SenseVoice 雲端 → SenseVoice 本地 → Whisper 本地）', '意圖分類（Intent）：高信心指令直接映射為動作，一般對話送入 LLM', 'AI 對話（LLM）：Qwen2.5-7B-Instruct 根據語境生成回覆', '語音合成（TTS）：edge-tts 雲端主線 + Piper 本地備援', '音訊播放：USB 外接喇叭或 Go2 Megaphone DataChannel']),
    ('2. 人臉辨識模組', ['YuNet 人臉偵測 + SFace 身份辨識 + IOU 多人追蹤', 'Hysteresis 三層穩定化機制', '辨識後觸發個人化 TTS 問候']),
    ('3. 手勢辨識模組', ['MediaPipe Gesture Recognizer 即時辨識', '手勢 → Go2 動作映射（如 open_palm → 停止）', '5 秒冷卻去重']),
    ('4. 姿勢辨識模組', ['MediaPipe Pose 33 關鍵點姿勢分類', '跌倒偵測 → 緊急警報', '持續狀態廣播']),
    ('5. 物體辨識模組', ['YOLO26n + TensorRT FP16 即時偵測 COCO 80 類', '高價值物體情境 TTS 話術', 'class_whitelist 動態篩選']),
    ('6. PawAI Studio 網頁控制台', ['即時感知面板（五大模組）', '文字與語音互動（push-to-talk）', '三欄即時影像串流', 'FastAPI + WebSocket Gateway 橋接 ROS2']),
    ('7. 統一中控模組（Interaction Executive）', ['State Machine 狀態管理', '優先序仲裁：EMERGENCY > stop > speech > gesture > face', '事件去重與冷卻期', '動作分派至 TTS 和 Go2']),
]:
    para(title, bold=True)
    for item in items:
        bullet(item)

# ── 四、背景知識 ──
h2('四、背景知識')
for title, content in [
    ('(一) 機器人作業系統（ROS2）', 'ROS2 為機器人軟體開發的彈性框架，採分散式架構，各功能模組封裝為獨立節點（Nodes），透過 Publisher/Subscriber 模型非同步交換資料。本系統利用此特性開發統一中控模組。'),
    ('(二) Unitree Go2 四足機器人平台', 'Go2 Pro 售價約 2,800 美元，配備 Hesai XT16 LiDAR。本專案透過 WebRTC DataChannel 控制。相較 Go2 EDU 版本可透過 CycloneDDS 直連，Pro 版本的 WebRTC 路徑延遲較高但成本較低。'),
    ('(三) MediaPipe 框架', 'Google 開發的即時多模態感知框架。本系統使用 Gesture Recognizer（手勢辨識，7 種內建手勢）和 Pose Estimation（姿勢辨識，33 個關鍵點），純 CPU 運行即可達到即時推理速度。'),
    ('(四) YuNet + SFace 人臉技術', 'OpenCV 內建的 YuNet 人臉偵測器（2023 年版）與 SFace 人臉辨識器（2021 年版）。YuNet 在 Jetson CPU 上可達即時推理速度。SFace 產出 128 維特徵向量進行餘弦相似度比對。'),
    ('(五) NVIDIA Jetson Orin Nano', '專為 AI 邊緣推理設計的嵌入式運算模組，8GB 統一記憶體，搭配 JetPack 6。負責運行所有本地感知模組和 ROS2 runtime。'),
    ('(六) Intel RealSense D435', 'RGB-D 深度攝影機，提供 RGB 影像與紅外線結構光深度感測。用於人臉偵測、手勢辨識、姿勢辨識與物體偵測。'),
    ('(七) SenseVoice / Whisper 語音辨識', '三級 ASR：SenseVoice 雲端（FunASR on RTX 8000，~600ms）→ SenseVoice 本地（sherpa-onnx int8 on Jetson CPU）→ Whisper 本地（faster-whisper CUDA on Jetson）。'),
    ('(八) edge-tts / Piper 語音合成', '雙重 TTS：edge-tts 雲端（微軟 Edge TTS，低延遲高音質）為主線；Piper 本地（ONNX 離線模型）為備援。'),
    ('(九) Qwen2.5-7B-Instruct 大語言模型', '阿里雲開源 LLM，透過 vLLM 部署於 RTX 8000。生成 ≤25 字中文回覆。雲端不可用時降級為 RuleBrain 規則引擎。'),
    ('(十) YOLO26n 物體辨識', 'YOLO26 Nano 版本（9.5MB），ONNX Runtime + TensorRT FP16，偵測 COCO 80 類物體。'),
    ('(十一) WebRTC DataChannel', '透過 WebRTC DataChannel 與 Go2 Pro 雙向通訊，包含動作控制指令與音訊播放（Megaphone API）。'),
    ('(十二) Next.js + FastAPI', '前端：Next.js 16 + React 19 + TypeScript + Tailwind CSS。後端：FastAPI + rclpy。WebSocket 即時轉發感知事件。'),
]:
    para(title, bold=True)
    para(content)

# ── 五、系統限制 ──
h2('五、系統限制（可行性分析）')
para('本節描述使用者可能期望但未包含在系統範圍內的需求。')
for title, importance, reason in [
    ('1. 自主導航避障', '使用者期望機器狗能自動避開障礙物、走到指定位置。重要性：高。', '未包含原因：D435 攝影機安裝角度偏上方，低障礙物直到極近距離才進入視野；加上偵測到停止的延遲鏈較長，三輪真機防撞測試均未通過。LiDAR 覆蓋率有限。程式碼完成但硬體限制無法克服，已停用。'),
    ('2. 機器狗本體語音對話', '使用者期望直接對機器狗說話，不需額外裝置。重要性：高。', '未包含原因：Go2 內建散熱風扇持續噪音，USB 麥克風訊噪比不足，實測語音辨識正確率顯著低於可用門檻。改以 PawAI Studio 網頁端 push-to-talk 作為語音入口。'),
    ('3. 多人同時互動', '使用者期望多人同時互動。重要性：中。', '未包含原因：人臉可追蹤最多 5 人，但手勢/姿勢辨識僅支援單人，且缺乏跨模組 ID 關聯機制。'),
    ('4. 多語言支援', '使用者期望支援英文、台語等。重要性：中。', '未包含原因：ASR/TTS/LLM 均針對中文調校，時程限制下優先確保中文穩定。'),
    ('5. 自動充電', '使用者期望機器狗自動回充電座。重要性：中。', '未包含原因：Go2 Pro 無自動對接充電座硬體；Jetson 供電依賴降壓模組，偶發不穩。'),
    ('6. 物體互動指令', '使用者期望「幫我拿水杯」。重要性：低。', '未包含原因：Go2 Pro 無機械手臂。物體辨識僅做偵測+情境 TTS。'),
]:
    para(title, bold=True)
    para(importance)
    para(reason)

# ══════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════
page_break()
h1('第二章 軟體需求規格')
h2('一、功能需求')

para('使用者角色說明', bold=True)
add_data_table(['角色', '說明'], [
    ['互動使用者', '與機器狗直接互動的人（長者、家屬、訪客），透過語音、手勢、臉部表情與機器狗互動。'],
    ['操作者', '負責監控與操控系統的人（開發人員、展示人員、教授），透過 PawAI Studio 網頁遠端操作。'],
])

para('使用者故事對應', bold=True)
bullet('優先順序：1: 第一階段、2: 第二階段')
add_data_table(['Epic', 'User Story', '優先順序'], [
    ['1. 透過語音與機器狗互動', '1.1 穩定接收音訊輸入', '1'],
    ['', '1.2 將中文語音轉為文字（ASR）', '1'],
    ['', '1.3 映射固定語句至執行意圖', '1'],
    ['', '1.4 根據語境生成邏輯回覆（LLM）', '1'],
    ['', '1.5 轉換回覆文字為自然語音（TTS）', '1'],
    ['', '1.6 透過喇叭播放音訊', '1'],
    ['2. 辨識手勢控制機器狗', '2.1 [待填寫]', '1'],
    ['', '2.2 [待填寫]', '1'],
    ['', '2.3 [待填寫]', '2'],
    ['3. 辨識姿勢偵測異常', '3.1 即時狀態回饋', '1'],
    ['', '3.2 緊急跌倒告警', '1'],
    ['', '3.3 識別對象確認', '1'],
    ['', '3.4 模組運作感知', '1'],
    ['', '3.5 活動歷史回顧', '1'],
    ['4. 辨識人臉進行個人化互動', '4.1 即時偵測人臉', '1'],
    ['', '4.2 辨識已註冊人物身份', '1'],
    ['', '4.3 個人化問候互動', '1'],
    ['5. 透過網頁觀測與操控機器狗', '5.1 即時感知狀態監控', '1'],
    ['', '5.2 文字與語音互動', '1'],
    ['', '5.3 即時影像串流', '2'],
    ['6. 辨識物體觸發情境互動', '6.1 [待填寫]', '1'],
    ['', '6.2 [待填寫]', '2'],
])

# ── User Story Cards ──
para('使用者故事卡', bold=True)

# Epic 1
add_story_card('1. 透過語音與機器狗互動', 'Epic', '互動使用者',
    '我想要透過語音直接與機器狗進行指令與日常對話。',
    '能體驗自然且流暢的端到端人機互動閉環。',
    ['語音能被穩定接收，不需喊特定喚醒詞。(1.1)', '中文短句能被準確辨識轉為文字。(1.2)', '說出預設指令時能觸發動作意圖。(1.3)', '說出日常對話時能得到合理回覆。(1.4)', '回覆能被轉換成自然語音（edge-tts + Piper）。(1.5)', '語音回覆能從喇叭清楚播放。(1.6)'])

for num, name, need, value, ac in [
    ('1.1', '穩定接收音訊輸入', '在不需要喚醒詞的情況下，直接對系統說話。', '確保語音輸入穩定性，降低系統複雜度與出錯率。',
     ['Given：系統處於待命狀態，no-VAD 主線模式。', 'When：系統麥克風接收到外部音訊輸入。', 'Then：系統成功捕獲音訊並傳遞至 ASR 模組。']),
    ('1.2', '將中文語音轉為文字（ASR）', '系統能將中文短句準確轉換為文字。', '提供正確文字數據作為後續意圖判斷基礎。',
     ['Given：已成功接收含中文語音的音訊。', 'When：音訊送入 ASR（SenseVoice 雲端/本地、Whisper 本地三級降級）。', 'Then：ASR 輸出正確中文字串。']),
    ('1.3', '映射固定語句至執行意圖', '說出「坐下」時能對應到正確的動作意圖。', '確保常見展示指令做出準確且低延遲的反應。',
     ['Given：Intent Rule 接收到 ASR 轉換的中文字串。', 'When：字串符合預設關鍵字映射表。', 'Then：成功映射至可執行意圖標籤（Intent）。']),
    ('1.4', '根據語境生成邏輯回覆（LLM）', '對話時獲得針對非固定語句的動態回應。', '賦予機器狗邏輯思考能力，更自然的互動體驗。',
     ['Given：接收到無法被 Intent Rule 映射的對話文字。', 'When：文字傳送至 Qwen2.5-7B-Instruct。', 'Then：LLM 生成符合語境的簡短回覆。若不可用則降級為 RuleBrain。']),
    ('1.5', '轉換回覆文字為自然語音（TTS）', '機器狗能以語音回答，而非僅呈現文字。', '提供聽覺回饋，完成端到端語音互動閉環。',
     ['Given：系統已產出回覆文字。', 'When：文字送入 TTS（edge-tts 雲端 / Piper 本地）。', 'Then：TTS 成功合成自然流暢音訊。']),
    ('1.6', '透過喇叭播放音訊', '語音回覆能從喇叭發出。', '增強互動沉浸感與真實感。',
     ['Given：TTS 音訊已生成，Jetson 已透過 ROS2 連線至 Go2。', 'When：系統下達播放指令。', 'Then：音訊透過 USB 喇叭清晰播放。']),
]:
    add_story_card(f'{num} {name}', 'User Story', '互動使用者', need, value, ac, priority=1)

# Epic 2, 3 empty
add_story_card('2. 辨識手勢控制機器狗', 'Epic', '[待填寫]', '[待填寫]', '[待填寫]', ['[待填寫]'])
add_story_card('3. 辨識姿勢偵測異常', 'Epic', '[待填寫]', '[待填寫]', '[待填寫]', ['[待填寫]'])

# Epic 4 — Face (Roy)
add_story_card('4. 辨識人臉進行個人化互動', 'Epic', '互動使用者（長者、家屬、訪客）',
    '我想要讓機器狗認出我是誰，並根據我的身份做出個人化的回應。',
    '讓機器狗能像認識朋友一樣叫出名字打招呼，提升陪伴感與信任度。',
    ['有人走入鏡頭範圍時，系統能即時偵測到人臉。(4.1)', '偵測到的人臉能與已註冊使用者進行身份比對。(4.2)', '成功辨識後，機器狗能叫出名字打招呼。(4.3)'])

for num, name, need, value, ac in [
    ('4.1', '即時偵測人臉', '當我走到機器狗前方時，系統能即時發現我的存在。', '作為身份辨識與問候互動的前提，確保系統能感知到使用者的到來。',
     ['Given：D435 攝影機已啟動，face_identity_node 正在運行。', 'When：有人走入鏡頭前方約 1-3 公尺範圍內。', 'Then：系統透過 YuNet 偵測到人臉，出現新的追蹤項目（含 track_id 與邊界框），發布 track_started 事件。']),
    ('4.2', '辨識已註冊人物身份', '系統能認出我是誰，將我與人臉資料庫中已註冊的身份配對。', '讓後續互動能個人化，而非千篇一律地對待所有人。',
     ['Given：系統已偵測到人臉（4.1），且人臉資料庫中有已註冊使用者照片與 SFace 特徵向量。', 'When：偵測到的人臉經 SFace 比對達到相似度門檻且連續多幀穩定。', 'Then：追蹤項目的穩定名稱更新為該使用者名稱，發布 identity_stable 事件。若相似度不足則維持 unknown。']),
    ('4.3', '個人化問候互動', '當機器狗認出我時，能主動叫我的名字打招呼。', '讓使用者感受到機器狗「認識」自己，提升陪伴感與互動溫度。',
     ['Given：系統已辨識出已註冊使用者（4.2 的 identity_stable 已觸發）。', 'When：Interaction Executive 收到事件，且該使用者 30 秒內未被問候過。', 'Then：機器狗透過 TTS 播放「{名字}你好！」，不在冷卻期內重複問候同一人。']),
]:
    add_story_card(f'{num} {name}', 'User Story', '互動使用者', need, value, ac, priority=1)

# Epic 5 — Studio (Roy)
add_story_card('5. 透過網頁觀測與操控機器狗', 'Epic', '操作者（開發人員、展示人員、教授）',
    '我想要透過瀏覽器即時監看機器狗所有感知狀態，並能用文字或語音互動。',
    '不需要站在機器狗旁邊，就能遠端觀測、互動、診斷問題，提升展示與開發效率。',
    ['在 Studio 即時看到五大模組狀態。(5.1)', '用文字或語音與機器狗對話。(5.2)', '看到即時攝影機影像。(5.3)'])

for num, name, need, value, ac, prio in [
    ('5.1', '即時感知狀態監控', '在瀏覽器上即時看到五大模組的最新狀態。', '遠端掌握系統運作情況，大幅提升除錯與展示效率。',
     ['Given：Gateway 已連線至 ROS2，Studio 前端已透過 WebSocket 連線。', 'When：ROS2 感知模組發布新事件。', 'Then：對應面板數秒內更新（人臉：姓名+距離、手勢：類別+信心度、姿勢：狀態+警報、物體：類別、語音：階段+文字）。'], 1),
    ('5.2', '文字與語音互動', '在聊天面板上輸入文字或錄音，讓機器狗回應。', '不需要站在機器狗旁邊也能進行互動測試。',
     ['Scenario 1（文字）— Given：Studio 已連線。When：輸入文字並送出。Then：經 Intent → LLM → TTS，機器狗播出回覆，Studio 顯示 AI 文字。',
      'Scenario 2（語音）— Given：Studio 已連線且有麥克風權限。When：按住麥克風錄音並鬆開。Then：經 ASR → Intent → LLM → TTS，機器狗播出回覆。'], 1),
    ('5.3', '即時影像串流', '在 Live View 頁面看到即時影像與偵測結果疊加。', '遠端看到機器狗「眼中的世界」，確認感知模組正常運作。',
     ['Given：D435 啟動，各感知模組發布 debug_image，Gateway video bridge 運行中。', 'When：操作者開啟 /studio/live 頁面。', 'Then：三欄即時影像（人臉/姿勢手勢/物體），每欄顯示 FPS。'], 2),
]:
    add_story_card(f'{num} {name}', 'User Story', '操作者', need, value, ac, priority=prio)

# Epic 6 empty
add_story_card('6. 辨識物體觸發情境互動', 'Epic', '[待填寫]', '[待填寫]', '[待填寫]', ['[待填寫]'])

# ── Activity Diagrams ──
h2('三、活動圖說明使用者關係')
para('[待補 — 建議使用 draw.io 或 PlantUML 繪製]')

# ══════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════
page_break()
h1('第三章 軟體設計規格')
h2('一、資料庫設計')
para('本系統不使用傳統關聯式資料庫（RDBMS），資料以串流形式在 ROS2 節點之間傳遞。')
add_data_table(['資料類型', '儲存方式', '說明'], [
    ['人臉資料庫', '檔案系統 + pickle', '已註冊使用者人臉照片與 SFace 128 維特徵向量'],
    ['感知事件', 'ROS2 Topics（JSON）', '即時串流，不持久化'],
    ['前端狀態', 'Zustand in-memory', '瀏覽器端狀態管理，關閉即消失'],
])
para('人臉資料庫結構', bold=True)
para('儲存位置：/home/jetson/face_db/。每位使用者一個資料夾，內含 PNG 人臉照片。啟動時自動偵測、提取 SFace 特徵向量、計算中心向量並生成 pickle 模型（model_sface.pkl）。照片數量有變動時自動重新訓練。')

h2('二、介面設計')
para('介面藍圖一覽表', bold=True)
add_data_table(['編號', '名稱', '對應使用者故事'], [
    ['UI-01', 'Studio 首頁（Mission Control）', '5.1, 5.2'],
    ['UI-02', '聊天面板', '5.2'],
    ['UI-03', 'Live View 三欄影像', '5.3'],
    ['UI-04', '人臉面板', '5.1'],
    ['UI-05', '手勢面板', '5.1'],
    ['UI-06', '姿勢面板', '5.1'],
    ['UI-07', '物體面板', '5.1'],
    ['UI-08', '語音面板', '5.1'],
])
para('介面藍圖畫面', bold=True)
para('[截圖待補 — 請截取各頁面並附上畫面說明]')

h2('三、資源需求（預算經費）')
para('開發系統所需要的人力、軟體、硬體及對應的經費預估', bold=True)
add_data_table(['經費項目', '經費名稱', '金額', '說明'], [
    ['硬體需求', 'Unitree Go2 Pro 四足機器人', 'NT$90,000', '開發期間所需'],
    ['硬體需求', 'NVIDIA Jetson Orin Nano SUPER 8GB', 'NT$8,000', '開發期間所需'],
    ['硬體需求', 'Intel RealSense D435 RGB-D', 'NT$12,000', '開發期間所需'],
    ['硬體需求', 'USB 麥克風 + USB 喇叭', 'NT$600', '開發期間所需'],
    ['硬體需求', '遠端伺服器（5× RTX 8000）', 'NT$0', '學校提供'],
    ['軟體需求', 'Ubuntu 22.04 + ROS2 Humble', 'NT$0', 'Open Source'],
    ['軟體需求', 'Python 3.10 + Node.js 20', 'NT$0', 'Open Source'],
    ['軟體需求', 'vLLM + FastAPI + Next.js', 'NT$0', 'Open Source'],
    ['人力需求', '5 人 × 預估 760 工時\n以時薪 NT$183 計', 'NT$139,080', '開發期間'],
    ['合計', '', 'NT$249,680', ''],
])

para('營運系統所需要的人力、軟體、硬體及對應的前三年經費預估', bold=True)
add_data_table(['經費項目', '經費名稱', '金額（每年）', '說明'], [
    ['硬體需求', '雲端 GPU 租賃（等同 1 張 A100）', 'NT$180,000', '若無學校伺服器'],
    ['硬體需求', 'Go2 + Jetson 硬體折舊（3 年）', 'NT$36,000', ''],
    ['軟體需求', 'Ubuntu + ROS2 + Open Source', 'NT$0', ''],
    ['人力需求', '維運 1 人兼職\n12 月 × NT$10,000', 'NT$120,000', '系統監控、模型更新'],
    ['其他', '網路、備品', 'NT$15,000', ''],
    ['年度合計', '', 'NT$351,000', ''],
    ['三年合計', '', 'NT$1,053,000', '沿用學校伺服器可降至 NT$513,000'],
])

# ══════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════
page_break()
h1('第四章 系統專題實作檢討')
h2('一、發展中遭遇到問題、困難與解決方法')

problems = [
    ('1. WebRTC 連線相容性問題', 'aiortc 套件自動升級至 1.14.0 後，SCTP 握手超時（30 秒以上），DataChannel 無法開啟，機器狗完全失聯。', '降級 aiortc 至 1.9.0 並移除 STUN 配置，SCTP 握手時間降至 0.5 秒。後續將版本鎖定寫入 requirements.txt。', '依賴套件版本必須鎖定，自動升級在嵌入式環境極易造成不可預期的相容性問題。'),
    ('2. Go2 Megaphone 音訊播放 API 格式陷阱', '團隊花費約兩週判定 Megaphone API「已失效」，無法從機器狗喇叭播放語音。根因為 payload 格式錯誤 — chunk_size 必須為 4096 base64 字元、msg type 必須為 "req" 而非 "msg"，官方文件未記載。', '透過逆向分析 Go2 韌體通訊協議找出正確格式，修正後連續 20/20 測試通過。', '硬體 API 文件不完整時需逆向驗證，優先懷疑格式而非硬體。'),
    ('3. Go2 Pro 有線模式架構限制', '嘗試建立 Ethernet 控制通道以降低延遲，連線成功後卻完全沒有 ROS2 topic 資料流。', 'Go2 Pro 內部 Wi-Fi 與有線網段隔離，有線只能到達 MCU。放棄有線模式，維持 Wi-Fi WebRTC。', '硬體架構設計決定軟體的可能性邊界，應在投入開發前先確認硬體能力。'),
    ('4. SLAM 建圖與 Nav2 導航的多重挫折', '經歷 Nav2 參數不生效（未 colcon build）、/scan 全為 inf（intensity 過濾過嚴）、Jetson 點雲頻率掉至 2 Hz（多實例 driver + 低效序列化，經七輪優化恢復至 5.1 Hz）、D435 鏡頭角度導致低障礙物看不到（三輪真機防撞測試全部撞上）等問題。', '導航避障功能完全停用。程式碼保留（含 20 個 unit tests），Demo 改為人工監控 + 手勢 stop。', '導航涉及硬體、網路、演算法三重限制，任一環節不達標整個功能即不可用。'),
    ('5. 多重網路跳轉導致控制延遲', '開發初期透過 Go2 (Wi-Fi) → Mac VM → Windows 多層跳轉進行 Nav2 控制，延遲 20-50ms 接近控制週期，指令積壓導致機器狗原地旋轉。', '改為 Windows 本地安裝 ROS2 + RViz2（零延遲控制），Foxglove 僅用於監控。', '即時控制應減少網路跳轉層數，選對工具比調參數更重要。'),
    ('6. 模型選型三次推翻', '手勢/姿勢辨識模型選型經歷三次推翻：RTMPose（GPU 滿載）→ DWPose（TRT 精度暴跌 + MMPose JetPack 6 零支援）→ MediaPipe（原判定不可行，實測推翻，CPU 即時推理）。', '最終採用 MediaPipe CPU-only 方案，手勢 7.2 FPS、姿勢 18.5 FPS，GPU 佔用 0%。', '文獻調查不能替代實測，Jetson 生態碎片化嚴重。'),
    ('7. ASR 三級 Fallback 的逐步演進', 'Whisper Small 中文短句+噪音辨識率已達瓶頸，無法滿足 Demo 需求。', '引入 SenseVoice 雲端（辨識率大幅提升）+ SenseVoice 本地（離線備援）形成三級降級。另實作 10+ 項 workaround（幻覺過濾、Echo Gate 等）。', '語音辨識在真實噪音環境下表現遠低於實驗室數據，需多層防禦。'),
    ('8. 跌倒偵測正面站姿誤判', '使用者正面站在鏡頭前時被錯誤判定為「跌倒」（肩膀展開導致 bbox 寬高比觸發條件）。', '新增 vertical_ratio 防護條件（肩膀到髖部垂直距離 / 軀幹長度），相對尺度不受距離影響。', '基於 bbox 的簡單規則在邊界情況易失效，需多維度防護。'),
    ('9. Jetson 環境污染事件', '在 Jetson 上執行 pip install ultralytics，自動拉入不相容的 PyTorch + NumPy 版本，整個 Python 環境崩潰。', '環境修復耗時半天。建立規範：Jetson 嚴禁直接 pip install 大型框架，正確做法是開發機匯出 ONNX 再 scp 到 Jetson。', '邊緣裝置 Python 環境極脆弱，任何可能覆蓋系統 wheel 的安裝都是高風險。'),
    ('10. LLM 視覺避障的根本局限', '早期嘗試透過 LLM 分析攝影機影像進行避障。LLM 能識別障礙物且建議轉向正確，但無法判斷「機器狗自身大小，能否通過」。', '評估多個 LLM 模型後認定純視覺方案只適合粗略避障。精確導航仍需 SLAM + 深度感測。', 'LLM 視覺理解有根本局限，不能完全取代傳統感測器。'),
    ('11. ROS2 QoS 策略不匹配', '多個模組間因 QoS 策略不一致導致訊息無法接收（RELIABLE subscriber 無法接收 BEST_EFFORT publisher）。', '統一感測器類 topic 使用 BEST_EFFORT QoS。', '多模組開發時 QoS 策略需在介面契約中明確規範。'),
    ('12. Jetson 供電危機', 'Sprint 11 天期間 Jetson 反覆被強制關機（Go2 電池 → XL4015 降壓 → Jetson，高負載 spike 超出降壓模組能力）。', '調整電壓 18.8V → 19.2V 稍有改善但未根治。至 Sprint 結束時仍為最大硬體風險。', '嵌入式系統供電設計需考慮負載峰值而非僅平均功耗。'),
    ('13. Go2 風扇噪音導致機身 ASR 不可用', '在 Go2 真機上 ASR 正確率因散熱風扇寬頻噪音驟降至不可用水準。調高 mic_gain 無效（噪音同步放大）。', '策略轉向：Demo 改為「視覺互動為主 + 網頁語音輔助」，語音入口從機身移到 PawAI Studio。', '物理環境限制無法透過軟體增益解決，需從系統架構層面尋找替代方案。'),
    ('14. SSH Tunnel 與遠端 GPU Server 連線', '雲端 ASR/LLM 依賴 SSH tunnel 連接 RTX 8000，tunnel 不穩定導致全 timeout。', '推 async fix 到 Server 並重啟；Jetson 建立 systemd user service 永久化 SSH tunnel。', '遠端連線涉及多層級，需逐層驗證每個環節。'),
]

for title, problem, solution, lesson in problems:
    para(title, bold=True)
    para(f'問題：{problem}')
    para(f'解決方法：{solution}')
    para(f'教訓：{lesson}')

h2('二、系統優缺點（SWOT）評估')
add_data_table(['', '正面', '負面'], [
    ['內部因素', '優勢(S)：六大模組完成度高；三級降級策略；模組化 ROS2 架構可獨立替換；PawAI Studio 遠端觀測能力', '劣勢(W)：供電不穩；機身 ASR 不可用；人臉 track 抖動未根治；單人追蹤限制'],
    ['外部因素', '機會(O)：長者陪伴市場增長；邊緣 AI 硬體降價；ROS2 生態成熟', '威脅(T)：Jetson 生態碎片化；Go2 Pro 韌體不開放；競品資源差距大'],
])

h2('三、發展心得')
para('[待團隊各成員填寫個人心得]')

h2('四、未來展望')
for item in [
    '導航避障 → 重新設計 D435 安裝角度（俯視 -45°）或升級 Go2 EDU（CycloneDDS 直連）',
    '機身語音 → 指向性麥克風 + 軟體降噪（RNNoise）',
    '多人互動 → 跨模組 track ID 關聯（face track_id 綁定 gesture/pose）',
    '物體互動 → 結合機械手臂或語音引導使用者操作',
    '供電穩定化 → 獨立 18650 電池組或升級降壓模組',
    'AI 大腦進化 → LLM-driven 決策中樞取代 rule-based Executive',
]:
    bullet(item)

# ── Appendix ──
page_break()
h1('附錄')
h2('一、文件分工及貢獻度說明')
para('[待團隊填寫 — 請列出文件章節並標示參與者之實際工時]')

h2('二、程式分工及貢獻度說明')
para('[待團隊填寫 — 請列出系統程式分工項目並標示參與者之實際工時]')
add_data_table(['負責人員 / 工作項目', '盧柏宇', '楊沛蓁', '鄔雨彤', '陳如恩', '黃旭'], [
    ['系統架構設計', '[工時]', '', '', '', ''],
    ['語音模組', '[工時]', '', '', '', ''],
    ['人臉辨識', '[工時]', '', '', '', ''],
    ['手勢辨識', '', '', '', '[工時]', '[工時]'],
    ['姿勢辨識', '', '', '', '[工時]', '[工時]'],
    ['物體辨識', '[工時]', '', '', '', ''],
    ['Interaction Executive', '[工時]', '', '', '', ''],
    ['PawAI Studio Gateway', '[工時]', '', '', '', ''],
    ['PawAI Studio 前端', '', '', '[工時]', '', ''],
    ['整合測試', '[工時]', '', '', '', ''],
    ['專題文件', '[工時]', '[工時]', '[工時]', '[工時]', '[工時]'],
    ['貢獻度', '[%]', '[%]', '[%]', '[%]', '[%]'],
])

h2('三、參考資料')
for i, ref in enumerate([
    'Unitree Robotics 文檔中心：https://support.unitree.com/home/zh/developer/about_Go2',
    'ROS2 Humble 官方文件：https://docs.ros.org/en/humble/',
    'MediaPipe 官方文件：https://ai.google.dev/edge/mediapipe/solutions/guide',
    'OpenCV Face Detection (YuNet)：https://docs.opencv.org/4.x/d0/dd4/tutorial_dnn_face.html',
    'SenseVoice 語音辨識：https://github.com/FunAudioLLM/SenseVoice',
    'faster-whisper：https://github.com/SYSTRAN/faster-whisper',
    'edge-tts：https://github.com/rany2/edge-tts',
    'Piper TTS：https://github.com/rhasspy/piper',
    'Qwen2.5 大型語言模型：https://github.com/QwenLM/Qwen2.5',
    'vLLM 推理框架：https://github.com/vllm-project/vllm',
    'YOLO (Ultralytics)：https://github.com/ultralytics/ultralytics',
    'Intel RealSense D435：https://www.intelrealsense.com/depth-camera-d435/',
    'NVIDIA Jetson Orin Nano：https://developer.nvidia.com/embedded/jetson-orin-nano',
    'Next.js：https://nextjs.org/',
    'FastAPI：https://fastapi.tiangolo.com/',
    'WebRTC：https://webrtc.org/',
], 1):
    para(f'{i}. {ref}')

# ── Save ──
output = '/home/roy422/newLife/elder_and_dog/docs/thesis/114-thesis.docx'
doc.save(output)
print(f'Saved: {output}')
