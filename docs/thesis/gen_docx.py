"""Generate 114-thesis.docx matching the original PDF format."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Style setup ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微軟正黑體'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微軟正黑體'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

# ── Helpers ──────────────────────────────────────────────
def add_page_break():
    doc.add_page_break()

def add_chapter_page(title):
    """Chapter title on its own page."""
    add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    add_page_break()

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table_2col(rows):
    """2-column table for user story cards. rows = [(label, value), ...]"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].width = Cm(4)
        table.rows[i].cells[1].text = value
        table.rows[i].cells[1].width = Cm(12)
        # Bold the label
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    doc.add_paragraph()  # spacing
    return table

def add_data_table(headers, rows):
    """Generic table with headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = str(val)
    doc.add_paragraph()
    return table

def add_user_story_card(name, category, role, need, value, acceptance, priority=None):
    """User Story card in 2-column table format matching original PDF."""
    rows = [
        ('名稱', name),
        ('類別', category),
    ]
    # Role/Need/Value as merged content
    table = doc.add_table(rows=2 + (1 if priority else 0), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Name
    table.rows[0].cells[0].text = '名稱'
    table.rows[0].cells[1].text = name
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs: r.bold = True

    # Row 1: Category
    table.rows[1].cells[0].text = '類別'
    table.rows[1].cells[1].text = category
    for p in table.rows[1].cells[0].paragraphs:
        for r in p.runs: r.bold = True

    # Priority row
    if priority:
        table.rows[2].cells[0].text = '優先順序'
        table.rows[2].cells[1].text = str(priority)
        for p in table.rows[2].cells[0].paragraphs:
            for r in p.runs: r.bold = True

    doc.add_paragraph()

    # Role/Need/Value as paragraph text below table (matching PDF format)
    p = doc.add_paragraph()
    run = p.add_run(f'角色：{role}')
    p = doc.add_paragraph()
    run = p.add_run(f'需求：{need}')
    p = doc.add_paragraph()
    run = p.add_run(f'價值：{value}')
    doc.add_paragraph()

    # Acceptance criteria
    add_para('接受條件：', bold=True)
    for item in acceptance:
        add_bullet(item)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('114專題文件')
run.font.size = Pt(24)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('Paw AI 機器狗')
run.font.size = Pt(28)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
p.add_run('指導老師：董惟鳳\n')
p.add_run('412401123 楊沛蓁\n')
p.add_run('412401355 盧柏宇\n')
p.add_run('412401135 鄔雨彤\n')
p.add_run('412401082 陳如恩\n')
p.add_run('412401094 黃  旭')

# ══════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════
add_chapter_page('第一章 系統描述')

add_h1('一、發展背景與動機')

add_h2('研究背景')
add_para(
    '近年來，人工智慧技術發展迅速，特別是大型語言模型（Large Language Models, LLMs，如 Qwen、LLaMA）'
    '的突破，使得機器系統具備了深度的語意理解與邏輯推理能力。同時，邊緣運算與電腦視覺'
    '（如 MediaPipe 姿態辨識、YuNet 人臉偵測技術）以及語音處理技術（如 SenseVoice 語音辨識與 '
    'edge-tts 語音合成）的日益成熟，讓資訊系統能夠以多模態（Multi-modal）的方式，精準感知周遭環境與人類的真實意圖。'
)
add_para(
    '在終端硬體載具方面，四足機器人（如仿生機器狗）的機電控制與運動靈活性已有顯著的提升。'
    '然而，綜觀目前市面上的實體機器人應用，多半仍依賴傳統的遙控設備、手機 App 或是僵化的預設腳本進行操作。'
    '現有系統缺乏將「雲端高階 AI 大腦」與「邊緣終端硬體（ROS2）」深度結合的解決方案，'
    '導致人機互動往往流於單向、死板的指令接收，機器無法根據人類自然的肢體語言（手勢）、'
    '語音指令或當下的環境脈絡，進行即時的動態感知與自主決策。'
)

add_h2('研究動機')
add_para(
    '基於上述背景，本專題「PawAI」旨在打破傳統實體機器人僵化且高門檻的互動模式。'
    '我們運用資訊管理領域的系統整合（System Integration）與軟體工程技術，'
    '開發一套無縫串接視覺感知、語音互動與大型語言模型決策的智能控制系統。'
)
add_para('本研究的核心動機包含以下幾點：')

add_para(
    '1. 實現自然且直覺的人機互動：透過建構完整的多模態感知管線（Multi-modal Perception Pipeline），'
    '讓系統能即時捕捉並解析使用者的手勢與臉部特徵；結合語音雙向處理，'
    '讓使用者能以最自然的「說話」與「動作」與機器狗互動，免除繁瑣的傳統控制介面。'
)
add_para(
    '2. 賦予終端設備動態決策能力：藉由串接大型語言模型 Qwen2.5-7B-Instruct'
    '（透過 vLLM 部署於遠端 RTX 8000 伺服器）作為系統的核心邏輯中樞，'
    '將接收到的多模態感知資料轉化為結構化的意圖（Intent），'
    '讓機器狗不再只是單純執行寫死的程式碼，而是能「理解」複雜情境並規劃後續行為。'
)
add_para(
    '3. 軟硬體整合與控制迴路優化：透過建構穩定的 ROS2 節點通訊與統一中控模組'
    '（Interaction Executive），確保 AI 模型的決策結果能以極低的延遲、安全且精確地'
    '轉化為實體機器狗的具體動作，完成從「感知」、「思考」到「行動」的完整閉環。'
)
add_para(
    '總結而言，本專題期望透過 PawAI 系統的建置，探索大型語言模型結合實體機器人的無限可能，'
    '不僅為智慧陪伴、巡檢或展演等應用場景提供具高度擴充性的架構範例，'
    '更進一步實現真正意義上的智慧人機共作（Human-Robot Collaboration）。'
)

# ── 二、系統發展目的 ──
add_h1('二、系統發展目的')

add_data_table(
    ['現有問題', '本系統之解決方法'],
    [
        ['實體機器人多依賴遙控器或手機 App 操作，互動方式單向且僵化。',
         '整合 MediaPipe 手勢/姿勢辨識與語音模組，實現自然直覺的多模態（手勢、語音、人臉）互動。'],
        ['機器人僅能執行預先寫死的腳本（Hard-coding），缺乏情境理解與自主決策能力。',
         '串接大語言模型（Qwen2.5-7B-Instruct）作為核心大腦，依據當下脈絡動態生成行為回饋。'],
        ['雲端 AI 模型與邊緣硬體設備之間，缺乏穩定、低延遲的整合與執行機制。',
         '開發基於 ROS2 的統一中控模組（Interaction Executive），精準且安全地驅動終端硬體作動。'],
        ['傳統機器人系統架構較為封閉，難以輕易擴充新功能或進行系統除錯。',
         '建立模組化感知管線與 PawAI Studio 網頁控制台，確保系統具備高擴充性與可觀測性。'],
    ]
)

add_para('(一) 整合多模態感知技術，提供自然直覺的人機互動體驗', bold=True)
add_para(
    '打破傳統單向且受限的控制介面，本系統導入 MediaPipe 模型進行即時的姿態與手勢推論，'
    '並輔以 YuNet + SFace 人臉辨識技術鎖定互動對象。同時，結合 STT（語音轉文字）與 TTS（文字轉語音）模組，'
    '讓使用者能以最日常的肢體動作與口語對話，與系統進行雙向、流暢的互動。'
)

add_para('(二) 串接大語言模型，建構具備情境理解的自主決策中樞', bold=True)
add_para(
    '為賦予機器狗動態決策的能力，本專案將前端擷取的多模態感知資料轉化為結構化的意圖（Intent），'
    '並透過 HTTP API 串接部署於遠端伺服器的大語言模型（Qwen2.5-7B-Instruct）。'
    '藉此，系統能深入理解使用者的語意與當下的互動脈絡，取代傳統預先寫死的控制腳本，'
    '自主規劃並生成最合適的行為回饋。'
)

add_para('(三) 開發 ROS2 統一中控，實現低延遲且穩定的硬體作動閉環', bold=True)
add_para(
    '為確保雲端 AI 的決策能精確落實於邊緣實體設備，本系統建構了基於 ROS2 的節點通訊機制。'
    '透過統一中控模組 Interaction Executive（以 State Machine 實現事件路由與優先序仲裁），'
    '系統能將抽象的行為指令，以極低的延遲轉譯為機器狗能直接執行的動作控制訊號，'
    '達成從「感知」、「運算」到「作動」的無縫銜接。'
)

add_para('(四) 建立 PawAI Studio 網頁控制台，確保系統的高可觀測性與可維護性', bold=True)
add_para(
    '導入軟體工程的模組化設計思維，建立 PawAI Studio 作為系統的統一觀測與操控入口。'
    'Studio 整合即時感知面板（人臉、手勢、姿勢、物體、語音狀態）、語音/文字互動介面、'
    '以及即時影像串流，方便團隊進行後續的效能監控與錯誤排查，更保留了未來擴增功能的彈性。'
)

add_para('非功能性需求', bold=True)
add_bullet('低延遲：語音端到端（E2E）互動延遲目標 < 5 秒（含 ASR + LLM + TTS）')
add_bullet('高可用性：三級降級策略（雲端 → 本地 → 規則），確保無網路時仍可基本互動')
add_bullet('模組化：每個感知模組為獨立 ROS2 package，可單獨替換或升級')
add_bullet('安全性：統一中控模組仲裁所有動作指令，緊急事件（跌倒）具最高優先序')

# ── 三、系統範圍 ──
add_h1('三、系統範圍')
add_para('本節根據使用者的互動情境，盤點系統的功能性需求。以下各功能模組構成 PawAI 系統的完整範圍，讓讀者在閱讀第二章需求規格前，對系統功能有概略性的了解。')

modules = [
    ('(一) 語音互動模組', [
        '語音辨識（ASR）：將使用者的中文語音即時轉換為文字，採用三級自動降級機制（SenseVoice 雲端 → SenseVoice 本地 → Whisper 本地）',
        '意圖分類（Intent）：對辨識出的文字進行意圖分析，高信心指令直接映射為動作，一般對話送入 LLM',
        'AI 對話（LLM）：透過大語言模型（Qwen2.5-7B-Instruct）根據對話語境生成簡短合理的回覆',
        '語音合成（TTS）：將回覆文字合成為自然語音，採用 edge-tts 雲端合成為主、Piper 本地合成為備援',
        '音訊播放：透過 USB 外接喇叭或 Go2 Megaphone DataChannel 播放語音回覆',
    ]),
    ('(二) 人臉辨識模組', [
        '人臉偵測：使用 YuNet 模型即時偵測 RGB 影像中的人臉位置',
        '身份辨識：使用 SFace 模型提取人臉特徵向量，與資料庫中已註冊的使用者進行比對',
        '多人追蹤：透過 IOU 演算法追蹤多張人臉（最多 5 人），各自獨立進行身份辨識',
        '身份穩定化：透過 Hysteresis 三層穩定機制，避免單幀抖動導致的身份誤判',
        '問候互動：辨識出已註冊使用者後，觸發 Interaction Executive 發送個人化 TTS 問候',
    ]),
    ('(三) 手勢辨識模組', [
        '手勢偵測與分類：使用 MediaPipe Gesture Recognizer 即時辨識手勢類別（如 thumbs_up、open_palm、stop 等）',
        '動作映射：將辨識出的手勢轉譯為對應的機器狗動作指令',
        '去重機制：同一手勢在 5 秒冷卻期內不重複觸發',
    ]),
    ('(四) 姿勢辨識模組', [
        '姿勢分類：使用 MediaPipe Pose 辨識人體 33 個關鍵點，分類為 standing、sitting、crouching、lying、fallen 等狀態',
        '跌倒偵測：當偵測到 fallen 狀態時，觸發緊急警報（EMERGENCY），機器狗發出語音提醒',
        '持續狀態廣播：以事件（Event）形式發布姿勢變化，供 Interaction Executive 和 PawAI Studio 訂閱',
    ]),
    ('(五) 物體辨識模組', [
        '物體偵測：使用 YOLO26n 模型搭配 TensorRT FP16 加速，即時偵測 COCO 80 類物體',
        '情境互動：對高價值物體（如水杯、水瓶、書本）觸發 TTS 互動話術',
        '類別篩選：可透過 class_whitelist 參數動態調整偵測範圍',
    ]),
    ('(六) PawAI Studio 網頁控制台', [
        '即時感知面板：同時顯示人臉、手勢、姿勢、物體、語音五大模組的即時狀態',
        '文字與語音互動：透過聊天面板輸入文字指令，或按住麥克風錄音進行語音互動',
        '即時影像串流：三欄即時影像（人臉偵測 / 姿勢手勢 / 物體偵測畫面）',
        'Gateway 橋接：FastAPI + WebSocket 橋接 ROS2 感知事件至瀏覽器',
    ]),
    ('(七) 統一中控模組（Interaction Executive）', [
        '狀態機：管理 IDLE → GREETING → CONVERSING → EXECUTING → EMERGENCY 狀態轉移',
        '優先序仲裁：EMERGENCY（跌倒）> obstacle > stop > speech > gesture > face',
        '事件去重：同一事件在冷卻期內不重複處理',
        '動作分派：將決策結果透過 /tts 和 /webrtc_req topic 分別送至 TTS 和 Go2 動作執行',
    ]),
]

for title, items in modules:
    add_h2(title)
    for item in items:
        add_bullet(item)

# ── The rest of Chapter 1 is very long. For brevity, we'll add key sections ──

add_h1('四、背景知識')
# Keep existing background + add new ones
bg_sections = [
    ('(一) 機器人作業系統（ROS2）',
     'ROS2 採用分散式架構，各功能模組封裝為獨立節點（Nodes），透過 Publisher/Subscriber 模型非同步交換資料。本系統利用此特性開發統一中控模組 Interaction Executive。'),
    ('(二) Unitree Go2 四足機器人平台',
     'Go2 Pro 售價約 2,800 美元，配備 Hesai XT16 LiDAR。本專案透過 WebRTC DataChannel 控制。相較 Go2 EDU 可透過 CycloneDDS 直連，Pro 版本延遲較高但成本較低。'),
    ('(三) MediaPipe 框架',
     'Google 開發的即時多模態感知框架。本系統使用 Gesture Recognizer（手勢辨識）和 Pose Estimation（姿勢辨識），純 CPU 運行即可達到即時推理速度。'),
    ('(四) YuNet + SFace 人臉技術',
     'OpenCV 內建的 YuNet 人臉偵測器（2023 年版）與 SFace 人臉辨識器（2021 年版）。YuNet 在 Jetson CPU 上可達即時推理速度。SFace 產出 128 維特徵向量進行餘弦相似度比對。'),
    ('(五) NVIDIA Jetson Orin Nano',
     '專為 AI 邊緣推理設計的嵌入式運算模組，8GB 統一記憶體，搭配 JetPack 6。負責運行所有本地感知模組和 ROS2 runtime。'),
    ('(六) Intel RealSense D435',
     'RGB-D 深度攝影機，提供 RGB 影像與紅外線結構光深度感測。本系統用於人臉偵測、手勢辨識、姿勢辨識與物體偵測。'),
    ('(七) SenseVoice / Whisper 語音辨識',
     '三級 ASR 策略：SenseVoice 雲端（FunASR on RTX 8000）→ SenseVoice 本地（sherpa-onnx int8 on Jetson CPU）→ Whisper 本地（faster-whisper CUDA on Jetson）。'),
    ('(八) edge-tts / Piper 語音合成',
     '雙重 TTS 策略：edge-tts（雲端，微軟 Edge TTS 引擎，低延遲高音質）為主線；Piper（本地離線 ONNX 模型）為備援。'),
    ('(九) Qwen2.5-7B-Instruct 大語言模型',
     '阿里雲開發的開源 LLM，透過 vLLM 推理框架部署於遠端 RTX 8000 伺服器。生成簡短中文回覆（≤25字）。雲端不可用時降級為 RuleBrain 規則引擎。'),
    ('(十) YOLO26n 物體辨識',
     'YOLO26 Nano 版本（9.5MB），透過 ONNX Runtime + TensorRT FP16 加速，可偵測 COCO 80 類物體。'),
    ('(十一) WebRTC DataChannel',
     '本系統透過 WebRTC DataChannel 與 Go2 Pro 雙向通訊，包含動作控制指令與音訊播放（Megaphone API）。'),
    ('(十二) Next.js + FastAPI（PawAI Studio）',
     '前端：Next.js 16 + React 19 + TypeScript + Tailwind CSS。後端：FastAPI + rclpy（ROS2 Python Client）。WebSocket 即時轉發感知事件。'),
]

for title, content in bg_sections:
    add_h2(title)
    add_para(content)

# ── 五、系統限制 ──
add_h1('五、系統限制（可行性分析）')
add_para('本節描述使用者可能期望但未包含在系統範圍內的需求，說明其重要性以及未包含的原因。')

limitations = [
    ('1. 自主導航避障',
     '使用者期望機器狗能自動避開障礙物。重要性高。',
     '未包含原因：D435 攝影機安裝角度偏上方，低障礙物直到極近距離才進入視野；加上偵測到停止的延遲鏈較長，真機防撞測試均未通過。LiDAR 覆蓋率有限。程式碼完成但硬體限制無法克服，已停用。'),
    ('2. 機器狗本體語音對話',
     '使用者期望直接對機器狗說話。重要性高。',
     '未包含原因：Go2 內建散熱風扇持續噪音，USB 麥克風 SNR 嚴重不足，實測語音辨識正確率顯著低於可用門檻。改以 PawAI Studio 網頁端 push-to-talk 作為語音入口。'),
    ('3. 多人同時互動',
     '使用者期望多人同時互動。重要性中。',
     '未包含原因：人臉可追蹤最多 5 人，但手勢/姿勢辨識僅支援單人，且缺乏跨模組 ID 關聯機制。'),
    ('4. 多語言支援',
     '使用者期望支援英文、台語等。重要性中。',
     '未包含原因：ASR/TTS/LLM 均針對中文調校，擴充語言需逐一驗證，時程限制下優先確保中文穩定。'),
    ('5. 自動充電與長時間自主運行',
     '使用者期望機器狗自動回充電座。重要性中。',
     '未包含原因：Go2 Pro 無自動對接充電座硬體；Jetson 供電依賴降壓模組，高負載時偶發不穩。'),
    ('6. 物體互動指令',
     '使用者期望告訴機器狗「幫我拿水杯」。重要性低。',
     '未包含原因：物體辨識僅做偵測+情境 TTS，Go2 Pro 無機械手臂，物理上無法抓取物體。'),
]

for title, importance, reason in limitations:
    add_h2(title)
    add_para(importance)
    add_para(reason)

# ══════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════
add_chapter_page('第二章 軟體需求規格')

add_h1('一、功能需求')

add_h2('1. 使用者角色說明')
add_data_table(
    ['角色', '說明', '使用情境'],
    [
        ['互動使用者', '與機器狗直接互動的人（長者、家屬、訪客）', '站在機器狗前方，透過語音、手勢、臉部表情互動'],
        ['操作者', '負責監控與操控系統的人（開發人員、展示人員、教授）', '透過 PawAI Studio 網頁遠端操作'],
    ]
)

add_h2('2. 使用者故事對應表')
add_data_table(
    ['Epic', 'User Story', '優先級'],
    [
        ['1. 透過語音與機器狗互動', '1.1 穩定接收音訊輸入', '1'],
        ['', '1.2 將中文語音轉為文字（ASR）', '1'],
        ['', '1.3 映射固定語句至執行意圖（Intent Rule）', '1'],
        ['', '1.4 根據語境生成邏輯回覆（LLM）', '1'],
        ['', '1.5 轉換回覆文字為自然語音（TTS）', '1'],
        ['', '1.6 透過喇叭播放音訊', '1'],
        ['2. 辨識手勢控制機器狗', '2.1 [待填寫]', '1'],
        ['', '2.2 [待填寫]', '1'],
        ['', '2.3 [待填寫]', '2'],
        ['3. 辨識姿勢偵測異常', '3.1 即時狀態回饋', '1'],
        ['', '3.2 緊急跌倒告警', '1'],
        ['', '3.3 識別對象確認', '1'],
        ['', '3.4 模組運作感知', '1'],
        ['', '3.5 活動歷史回顧', '1'],
        ['4. 辨識人臉進行個人化互動', '4.1 即時偵測人臉', '1'],
        ['', '4.2 辨識已註冊人物身份', '1'],
        ['', '4.3 個人化問候互動', '1'],
        ['', '4.4 多人同時追蹤', '1'],
        ['5. 透過網頁觀測與操控機器狗', '5.1 即時感知狀態監控', '1'],
        ['', '5.2 文字與語音互動', '1'],
        ['', '5.3 即時影像串流', '2'],
        ['6. 辨識物體觸發情境互動', '6.1 [待填寫]', '1'],
        ['', '6.2 [待填寫]', '2'],
    ]
)

# ── User Story Cards ──
add_h1('二、使用者故事卡')

# Epic 1
add_user_story_card(
    '1. 透過語音與機器狗互動', 'Epic',
    '互動使用者',
    '我想要透過語音直接與機器狗進行指令與日常對話。',
    '能體驗自然且流暢的端到端（End-to-End）人機互動閉環。',
    [
        '能直接對機器狗說話，不需要喊特定的喚醒詞，語音即可被穩定接收。(1.1)',
        '中文短句能被系統準確辨識並轉換為文字。(1.2)',
        '說出預設指令時，系統能正確理解並觸發對應的動作意圖。(1.3)',
        '說出日常對話時，能得到根據對話情境所產生的簡短、合理回覆。(1.4)',
        '回覆能被轉換成自然且流暢的語音（edge-tts 雲端主線 + Piper 本地 fallback）。(1.5)',
        '語音回覆能從喇叭清楚播放出來。(1.6)',
    ]
)

# 1.1 - 1.6 (abbreviated — same content as PDF)
for us_num, us_name, us_need, us_value, us_ac in [
    ('1.1', '穩定接收音訊輸入', '在不需要額外喚醒詞的情況下，直接對系統說話。',
     '確保語音輸入流程的穩定性，降低 MVP 階段系統的複雜度與出錯率。',
     ['Given：系統處於待命狀態，且設定為 no-VAD 主線模式。',
      'When：系統麥克風接收到外部音訊輸入時。',
      'Then：系統應能成功捕獲該音訊，並穩定傳遞至下游 ASR 模組。']),
    ('1.2', '將中文語音轉為文字（ASR）', '系統能將我說出的中文短句準確轉換為文字。',
     '提供正確的文字數據，作為系統後續意圖判斷與大腦決策的基礎。',
     ['Given：系統已成功接收一段包含中文語音的音訊資料。',
      'When：系統將該音訊送入 ASR 模組（SenseVoice 雲端/本地、Whisper 本地三級降級）進行推論。',
      'Then：ASR 模組應在可接受延遲內輸出正確中文字串。']),
    ('1.3', '映射固定語句至執行意圖（Intent Rule）', '當我說出特定指令時，系統能對應到正確的動作意圖。',
     '確保機器狗能對常見展示指令做出準確且低延遲的直接動作反應。',
     ['Given：Intent Rule 模組接收到 ASR 轉換完成的中文字串。',
      'When：該字串符合系統預設的關鍵字或規則映射表。',
      'Then：系統成功映射至對應的可執行意圖標籤（Intent）。']),
    ('1.4', '根據語境生成邏輯回覆（LLM）', '在對話時獲得系統針對非固定語句的動態文字回應。',
     '賦予機器狗初步的邏輯思考能力，提供更自然的人機互動體驗。',
     ['Given：系統接收到無法被 Intent Rule 直接映射的日常對話文字。',
      'When：系統將該文字傳送至 Qwen2.5-7B-Instruct 進行運算。',
      'Then：LLM 生成符合語境且簡短合理的文字回覆。若 LLM 不可用，自動降級為 RuleBrain。']),
    ('1.5', '轉換回覆文字為自然語音（TTS）', '機器狗能以語音發聲方式回答我。',
     '提供直覺的聽覺回饋，完成端到端的完整語音互動閉環。',
     ['Given：系統已產出一段準備回覆的文字。',
      'When：文字送入 TTS 模組（edge-tts 雲端 / Piper 本地）處理。',
      'Then：TTS 成功合成自然流暢的音訊檔案。']),
    ('1.6', '透過喇叭播放音訊', '系統生成的語音回覆能從喇叭發出。',
     '增強互動的沉浸感與真實感。',
     ['Given：TTS 音訊已生成，Jetson 已透過 ROS2 連線至 Go2。',
      'When：系統下達播放音訊的控制指令。',
      'Then：音訊穩定透過 USB 外接喇叭清晰播放。']),
]:
    add_user_story_card(f'{us_num} {us_name}', 'User Story', '互動使用者', us_need, us_value, us_ac, priority=1)

# Epic 2, 3 — empty shells
for epic_num, epic_name in [('2', '辨識手勢控制機器狗'), ('3', '辨識姿勢偵測異常')]:
    add_user_story_card(f'{epic_num}. {epic_name}', 'Epic', '[待填寫]', '[待填寫]', '[待填寫]', ['[待填寫]'])

# Epic 4 — Face (Roy's)
add_user_story_card(
    '4. 辨識人臉進行個人化互動', 'Epic',
    '互動使用者（長者、家屬、訪客）',
    '我想要讓機器狗認出我是誰，並根據我的身份做出個人化的回應。',
    '讓機器狗的互動不再千篇一律，能像認識朋友一樣叫出名字打招呼，提升陪伴感與信任度。',
    [
        '有人走入鏡頭範圍時，系統能即時偵測到人臉。(4.1)',
        '偵測到的人臉能與已註冊使用者進行身份比對。(4.2)',
        '成功辨識後，機器狗能叫出名字打招呼。(4.3)',
        '多人同時出現時，系統能分別追蹤與辨識。(4.4)',
    ]
)

for us_num, us_name, us_need, us_value, us_ac in [
    ('4.1', '即時偵測人臉', '當我走到機器狗前方時，系統能即時發現我的存在。',
     '作為身份辨識與問候互動的前提，確保系統能感知到使用者的到來。',
     ['Given：D435 攝影機已啟動，face_identity_node 正在運行。',
      'When：有人走入鏡頭前方約 1-3 公尺範圍內。',
      'Then：系統偵測到人臉，出現新的 track（含 track_id 與 bounding box），發布 track_started 事件。']),
    ('4.2', '辨識已註冊人物身份', '系統能認出我是誰（比對已註冊的人臉資料庫）。',
     '讓後續互動能個人化，而非千篇一律地對待所有人。',
     ['Given：系統已偵測到人臉（4.1），且資料庫中有已註冊使用者。',
      'When：偵測到的人臉經 SFace 比對達到相似度門檻且連續穩定。',
      'Then：track 的 stable_name 更新為使用者名稱，發布 identity_stable 事件。']),
    ('4.3', '個人化問候互動', '當機器狗認出我時，能主動叫我的名字打招呼。',
     '讓使用者感受到機器狗「認識」自己，提升陪伴感與互動溫度。',
     ['Given：系統已辨識出已註冊使用者（4.2 的 identity_stable 已觸發）。',
      'When：Interaction Executive 收到事件，且該使用者 30 秒內未被問候過。',
      'Then：機器狗透過 TTS 播放「{名字}你好！」問候語，不在冷卻期內重複。']),
    ('4.4', '多人同時追蹤', '多人同時站在機器狗前方時，系統能分別追蹤每張臉。',
     '確保多人場景下系統不會混淆身份。',
     ['Given：鏡頭前方有 2 位以上使用者。',
      'When：系統同時偵測到多張人臉。',
      'Then：透過 IOU 追蹤各自獨立辨識（最多 5 人），track_id 不混淆。']),
]:
    add_user_story_card(f'{us_num} {us_name}', 'User Story', '互動使用者', us_need, us_value, us_ac, priority=1)

# Epic 5 — Studio (Roy's)
add_user_story_card(
    '5. 透過網頁觀測與操控機器狗', 'Epic',
    '操作者（開發人員、展示人員、教授）',
    '我想要透過瀏覽器即時監看機器狗的所有感知狀態，並能用文字或語音互動。',
    '不需要站在機器狗旁邊，就能遠端觀測、互動、診斷問題，提升展示與開發效率。',
    [
        '在 Studio 網頁上即時看到所有感知模組的狀態資料。(5.1)',
        '在 Studio 網頁上用文字或語音與機器狗對話。(5.2)',
        '在 Studio 網頁上看到即時攝影機影像。(5.3)',
    ]
)

for us_num, us_name, us_need, us_value, us_ac, prio in [
    ('5.1', '即時感知狀態監控', '在瀏覽器上即時看到五大模組的最新狀態。',
     '遠端即可掌握系統運作情況，大幅提升除錯與展示效率。',
     ['Given：Gateway 已連線至 ROS2，Studio 前端已透過 WebSocket 連線。',
      'When：ROS2 感知模組發布新事件。',
      'Then：對應面板數秒內更新（人臉：姓名+距離、手勢：類別+信心度、姿勢：狀態+警報、物體：類別、語音：階段+文字）。'], 1),
    ('5.2', '文字與語音互動', '在聊天面板上輸入文字或錄音，讓機器狗回應。',
     '不需要站在機器狗旁邊也能進行語音互動測試。',
     ['Scenario 1（文字）— Given：Studio 已連線 Gateway。When：操作者輸入文字並送出。Then：文字經 Intent → LLM → TTS，機器狗播出回覆，Studio 顯示 AI 文字。',
      'Scenario 2（語音）— Given：Studio 已連線，瀏覽器已取得麥克風權限。When：操作者按住麥克風錄音並鬆開。Then：錄音經 ASR → Intent → LLM → TTS，機器狗播出回覆。'], 1),
    ('5.3', '即時影像串流', '在 Live View 頁面看到即時影像與偵測結果。',
     '遠端看到機器狗「眼中的世界」，確認感知模組正常運作。',
     ['Given：D435 啟動，各感知模組發布 debug_image，Gateway video bridge 運行中。',
      'When：操作者開啟 /studio/live 頁面。',
      'Then：三欄即時影像（人臉/姿勢手勢/物體），每欄顯示 FPS。'], 2),
]:
    add_user_story_card(f'{us_num} {us_name}', 'User Story', '操作者', us_need, us_value, us_ac, priority=prio)

# Epic 6 — empty
add_user_story_card('6. 辨識物體觸發情境互動', 'Epic', '[待填寫]', '[待填寫]', '[待填寫]', ['[待填寫]'])

# ── Activity Diagrams ──
add_h1('三、活動圖說明使用者關係')
add_para('[待補 — 建議使用 draw.io 或 PlantUML 繪製活動圖]')

# ══════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════
add_chapter_page('第三章 軟體設計規格')

add_h1('一、資料庫設計')
add_para('本系統不使用傳統關聯式資料庫（RDBMS），資料以串流形式在 ROS2 節點之間傳遞。')

add_data_table(
    ['資料類型', '儲存方式', '說明'],
    [
        ['人臉資料庫', '檔案系統 + pickle', '已註冊使用者的人臉照片與特徵向量'],
        ['感知事件', 'ROS2 Topics（JSON）', '即時串流，不持久化'],
        ['前端狀態', 'Zustand in-memory', '瀏覽器端狀態管理，關閉即消失'],
    ]
)

add_h2('人臉資料庫結構（NoSQL — 檔案系統）')
add_para('儲存位置：/home/jetson/face_db/')
add_para('每位使用者一個資料夾，內含 PNG 人臉照片。系統啟動時自動偵測、提取 SFace 128 維特徵向量並生成 pickle 模型（model_sface.pkl）。')

add_h1('二、介面設計')
add_para('[截圖待補 — 以下為各介面的文字說明]')

add_data_table(
    ['編號', '介面名稱', '對應 User Story', '說明'],
    [
        ['UI-01', 'Studio 首頁（Mission Control）', '5.1, 5.2', '主控台：感知面板 + 聊天面板'],
        ['UI-02', '聊天面板', '5.2', '文字/語音輸入 + AI 回覆'],
        ['UI-03', 'Live View 三欄影像', '5.3', '即時攝影機影像 + 偵測結果'],
        ['UI-04', '人臉面板', '5.1', '追蹤卡片：姓名、相似度、距離'],
        ['UI-05', '手勢面板', '5.1', '手勢類別、信心度、歷史'],
        ['UI-06', '姿勢面板', '5.1', '姿勢狀態、跌倒警報'],
        ['UI-07', '物體面板', '5.1', '偵測物體：類別、信心度'],
        ['UI-08', '語音面板', '5.1', '語音管線階段、辨識文字'],
    ]
)

add_h1('三、資源需求')
add_h2('開發期間')
add_data_table(
    ['項目', '規格', '費用估算'],
    [
        ['Unitree Go2 Pro', '四足機器人', '~NT$90,000'],
        ['Jetson Orin Nano SUPER', '8GB 邊緣運算', '~NT$8,000'],
        ['Intel RealSense D435', 'RGB-D 攝影機', '~NT$12,000'],
        ['USB 麥克風 + 喇叭', '音訊 I/O', '~NT$600'],
        ['遠端伺服器（5× RTX 8000）', '雲端 LLM 推理', '學校提供'],
        ['合計', '', '~NT$110,800'],
    ]
)

# ══════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════
add_chapter_page('第四章 系統專題實作檢討')

add_h1('一、發展中遭遇問題')

problems = [
    ('1. 模型選型三次推翻',
     '原計畫 RTMPose wholebody → GPU 滿載棄用；改評估 DWPose → TensorRT 精度暴跌棄用；'
     'MediaPipe 原判定不可行 → 實測推翻，CPU 即可即時推理。教訓：文獻不能替代實測，Jetson 生態碎片化。'),
    ('2. Go2 Megaphone API 格式陷阱',
     '花費約兩週判定 API「失效」。根因是 payload 格式錯誤（chunk_size、msg type 必須為 "req"）。修正後 20/20 通過。教訓：硬體 API 文件不完整時需逆向驗證。'),
    ('3. Jetson 環境污染事件',
     'pip install ultralytics 破壞 Jetson 專用 PyTorch wheel，環境修復耗時半天。教訓：Jetson 嚴禁直接安裝可能覆蓋系統 wheel 的大型框架。'),
    ('4. 跌倒偵測正面站姿誤判',
     '肩膀展開導致 bbox 寬高比觸發跌倒判定。修復：新增 vertical_ratio guard（相對尺度，不受距離影響）。'),
    ('5. ASR 三級 Fallback 演進',
     'Whisper 中文短句+噪音辨識率達瓶頸 → 引入 SenseVoice 雲端大幅提升 → SenseVoice 本地成為可靠離線備援。額外實作 10+ workaround。'),
    ('6. 供電危機',
     'XL4015 降壓模組不足以支持高負載 spike。Sprint 期間多次斷電，至結束時仍為最大硬體風險。'),
]

for title, content in problems:
    add_h2(title)
    add_para(content)

add_h1('二、系統優缺點（SWOT）評估')
add_data_table(
    ['', '正面', '負面'],
    [
        ['內部因素',
         'S：六大模組完成度高；三級降級策略；模組化 ROS2 架構可獨立替換；PawAI Studio 遠端觀測',
         'W：供電不穩；機身 ASR 不可用；人臉 track 抖動未根治；單人追蹤限制'],
        ['外部因素',
         'O：長者陪伴市場增長；邊緣 AI 硬體降價；ROS2 生態成熟',
         'T：Jetson 生態碎片化；Go2 Pro 韌體不開放；競品資源差距大'],
    ]
)

add_h1('三、未來展望')
future = [
    '導航避障 → 重新設計 D435 安裝角度（俯視 -45°）或升級 Go2 EDU',
    '機身語音 → 指向性麥克風 + 軟體降噪（RNNoise）',
    '多人互動 → 跨模組 track ID 關聯（face track_id 綁定 gesture/pose）',
    '物體互動 → 結合機械手臂或語音引導使用者操作',
    '供電穩定化 → 獨立電池組或升級降壓模組',
    'AI 大腦進化 → LLM-driven 決策中樞取代 rule-based Executive',
]
for item in future:
    add_bullet(item)

add_h1('四、分工貢獻')
add_para('[待團隊填寫]')
add_data_table(
    ['成員', '主要負責模組', '預估工時', '貢獻度'],
    [
        ['盧柏宇', '系統架構、語音、人臉、物體、Executive、Studio Gateway、整合測試', '[待填]', '[待填]'],
        ['楊沛蓁', '[待填]', '[待填]', '[待填]'],
        ['鄔雨彤', 'PawAI Studio 前端 UI', '[待填]', '[待填]'],
        ['陳如恩', '手勢/姿勢辨識研究', '[待填]', '[待填]'],
        ['黃旭', '手勢/姿勢辨識研究', '[待填]', '[待填]'],
    ]
)

# ── References ──
add_page_break()
add_h1('參考資料')
refs = [
    'Unitree Robotics 文檔中心：https://support.unitree.com/home/zh/developer/about_Go2',
    'ROS2 Humble 官方文件：https://docs.ros.org/en/humble/',
    'MediaPipe 官方文件：https://ai.google.dev/edge/mediapipe/solutions/guide',
    'OpenCV Face Detection (YuNet)：https://docs.opencv.org/4.x/d0/dd4/tutorial_dnn_face.html',
    'SenseVoice 語音辨識：https://github.com/FunAudioLLM/SenseVoice',
    'faster-whisper：https://github.com/SYSTRAN/faster-whisper',
    'edge-tts：https://github.com/rany2/edge-tts',
    'Piper TTS：https://github.com/rhasspy/piper',
    'Qwen2.5 大型語言模型：https://github.com/QwenLM/Qwen2.5',
    'vLLM 推理框架：https://github.com/vllm-project/vllm',
    'YOLO (Ultralytics)：https://github.com/ultralytics/ultralytics',
    'Intel RealSense D435：https://www.intelrealsense.com/depth-camera-d435/',
    'NVIDIA Jetson Orin Nano：https://developer.nvidia.com/embedded/jetson-orin-nano',
    'Next.js：https://nextjs.org/',
    'FastAPI：https://fastapi.tiangolo.com/',
    'WebRTC：https://webrtc.org/',
]
for i, ref in enumerate(refs, 1):
    add_para(f'{i}. {ref}')

# ── Save ──
output_path = '/home/roy422/newLife/elder_and_dog/docs/thesis/114-thesis.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
print(f'Pages estimated: ~{len(doc.paragraphs) // 30 + len(doc.tables) * 2}')
